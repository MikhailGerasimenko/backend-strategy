"""Загрузка рассылки Kallanish из Word (.docx) для промпта брифа."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from .config import BRIEF_KALLANISH_MAX_CHARS

_KALLANISH_GLOB = ("*kallanish*.docx", "*Kallanish*.docx", "*KALLANISH*.docx")
KALLANISH_CANONICAL_NAME = "kallanish.docx"
MAX_KALLANISH_UPLOAD_BYTES = 30 * 1024 * 1024


def _is_ignorable_sidecar(path: Path) -> bool:
    """Файлы macOS (._*, .DS_Store) при копировании на Linux ломают чтение docx."""
    name = path.name
    return name.startswith("._") or name in {".DS_Store", ".gitkeep"}


def _is_valid_docx_candidate(path: Path) -> bool:
    if not path.is_file() or _is_ignorable_sidecar(path):
        return False
    if path.suffix.lower() != ".docx":
        return False
    try:
        return path.stat().st_size >= 100
    except OSError:
        return False


def discover_kallanish_docx(directory: Path) -> Path | None:
    """Ищет .docx с «kallanish» в имени в каталоге (обычно «Новости»)."""
    if not directory.is_dir():
        return None
    candidates: list[Path] = []
    for pattern in _KALLANISH_GLOB:
        for path in directory.glob(pattern):
            if _is_valid_docx_candidate(path):
                candidates.append(path)
    if not candidates:
        return None
    preferred = {name.lower() for name in ("kallanish.docx", "kallanish daily.docx")}
    for path in candidates:
        if path.name.lower() in preferred:
            return path
    return max(candidates, key=lambda item: item.stat().st_mtime)


def extract_docx_plain_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _truncate_text(text: str, max_chars: int) -> str:
    compact = re.sub(r"\n{3,}", "\n\n", text.strip())
    if len(compact) <= max_chars:
        return compact
    return compact[: max_chars - 1].rstrip() + "…"


def load_kallanish_text(path: Path, *, max_chars: int | None = None) -> str:
    limit = max_chars if max_chars is not None else BRIEF_KALLANISH_MAX_CHARS
    return _truncate_text(extract_docx_plain_text(path), limit)


def format_kallanish_for_prompt(path: Path, text: str) -> str:
    return f"Файл: {path.name}\n\n{text}"


def resolve_kallanish_path(
    *,
    explicit_path: Path | None,
    news_dir: Path | None,
    include: bool,
) -> Path | None:
    if not include:
        return None
    if explicit_path is not None:
        if explicit_path.exists() and _is_valid_docx_candidate(explicit_path):
            return explicit_path
        return None
    if news_dir is not None:
        return discover_kallanish_docx(news_dir)
    return None


def get_kallanish_info(directory: Path) -> dict[str, Any]:
    """Метаданные актуального файла Kallanish для API/веб-интерфейса."""
    path = discover_kallanish_docx(directory)
    if path is None:
        return {
            "has_file": False,
            "filename": None,
            "path": None,
            "size_kb": 0,
            "modified_at": None,
            "text_chars": 0,
        }
    stat = path.stat()
    try:
        text_len = len(extract_docx_plain_text(path))
    except Exception:  # noqa: BLE001
        text_len = 0
    return {
        "has_file": True,
        "filename": path.name,
        "path": str(path),
        "size_kb": round(stat.st_size / 1024, 1),
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "text_chars": text_len,
        "is_canonical": path.name.lower() == KALLANISH_CANONICAL_NAME.lower(),
    }


def save_kallanish_upload(
    directory: Path,
    content: bytes,
    original_filename: str,
) -> dict[str, Any]:
    """Сохраняет загруженный .docx как актуальный kallanish.docx."""
    name = (original_filename or "").strip()
    if not name.lower().endswith(".docx"):
        raise ValueError("Допустим только формат Word (.docx).")
    if len(content) > MAX_KALLANISH_UPLOAD_BYTES:
        raise ValueError(
            f"Файл слишком большой (макс. {MAX_KALLANISH_UPLOAD_BYTES // (1024 * 1024)} МБ)."
        )
    if len(content) < 100:
        raise ValueError("Файл пуст или повреждён.")

    directory.mkdir(parents=True, exist_ok=True)
    target = directory / KALLANISH_CANONICAL_NAME
    backup_name: str | None = None
    if target.exists():
        backup = directory / f"kallanish_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        target.rename(backup)
        backup_name = backup.name

    target.write_bytes(content)
    try:
        text = extract_docx_plain_text(target)
    except Exception as exc:  # noqa: BLE001
        target.unlink(missing_ok=True)
        if backup_name:
            (directory / backup_name).rename(target)
        raise ValueError(f"Не удалось прочитать Word-файл: {exc}") from exc

    if not text.strip():
        target.unlink(missing_ok=True)
        if backup_name:
            (directory / backup_name).rename(target)
        raise ValueError("В файле не найден текст (пустой документ).")

    info = get_kallanish_info(directory)
    info["backup_filename"] = backup_name
    info["uploaded_as"] = KALLANISH_CANONICAL_NAME
    return info


def build_kallanish_block(
    *,
    explicit_path: Path | None,
    news_dir: Path | None,
    include: bool,
) -> tuple[str, Path | None]:
    path = resolve_kallanish_path(
        explicit_path=explicit_path,
        news_dir=news_dir,
        include=include,
    )
    if path is None:
        if explicit_path is not None and not explicit_path.exists():
            return (
                f"(указанный файл Kallanish не найден: {explicit_path})",
                None,
            )
        return (
            "(файл Kallanish не найден — положите *kallanish*.docx в папку «Новости»)",
            None,
        )
    text = load_kallanish_text(path)
    if not text.strip():
        return (f"(файл {path.name} пуст или не удалось извлечь текст)", path)
    return format_kallanish_for_prompt(path, text), path
