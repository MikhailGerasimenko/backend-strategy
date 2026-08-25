"""Экспорт аналитического брифа в Word (.docx) напрямую из текста LLM."""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE_PATH = PROJECT_ROOT / "Шаблон брифа_v2.docx"
LOGO_CANDIDATES = (
    PROJECT_ROOT / "web/static/images/severstal-logo.png",
    Path(__file__).resolve().parent / "assets/severstal-logo-header.png",
)

TEMPLATE_PRODUCTS = (
    "Железная руда в Китае",
    "Коксующийся уголь",
    "Стальной лом в Турции",
    "ГК прокат",
)

TEMPLATE_RUSSIA_MARKET = (
    "Железная руда и окатыши в России",
    "Коксующийся уголь в России",
    "Стальной лом в России",
    "ГК прокат в России",
)

PRODUCT_KEYS = {
    "Железная руда в Китае": "ore",
    "Коксующийся уголь": "coal",
    "Стальной лом в Турции": "scrap",
}

TEMPLATE_COUNTRIES = (
    "Европейский Союз",
    "США",
    "Индия",
    "Китай",
    "Турция",
)

TEMPLATE_RUSSIA_TOPICS = (
    "Динамика ВВП России",
    "Динамика промышленного производства",
    "Жилищное строительство в России, ипотечное кредитование",
    "Динамика машиностроения в России: автомобилестроение, вагоностроение",
    "Динамика нефтегазового сектора в России: добыча нефти и газа",
    "Оценки настроений в промышленности, PMI",
    "Динамика инфляции в России",
    "Обменный курс рубля",
    "Ключевая ставка в России",
    "Федеральный бюджет",
    "Экономическая политика",
)

BALANCE_LABELS = (
    "производств",
    "экспорт",
    "потреблен",
)

NO_NEWS = "Релевантных новостей не выявлено."


def brief_docx_filename(news_date: date, brief_kind: str = "full") -> str:
    """Имя файла брифа: YYYYMMDD_<тип>_дайджест.docx (дата новостей)."""
    suffix_by_kind = {
        "full": "Новостной_дайджест",
        "market": "Рыночный_дайджест",
        "corporate": "Корпоративный_дайджест",
    }
    suffix = suffix_by_kind.get(brief_kind, suffix_by_kind["full"])
    return f"{news_date.strftime('%Y%m%d')}_{suffix}.docx"


def brief_docx_filenames_for_day(news_date: date) -> list[str]:
    return [brief_docx_filename(news_date, kind) for kind in ("full", "market", "corporate")]


PERIOD_BRIEF_KIND_TITLES: dict[str, str] = {
    "full": "Полный бриф",
    "market": "Рыночный бриф",
    "corporate": "Новостной бриф",
    "monthly": "Полный ежемесячный бриф",
    "monthly_news": "Ежемесячный новостной бриф",
    "monthly_market": "Ежемесячный рыночный бриф",
    "monthly_corporate": "Ежемесячный рыночный бриф",
}


def _sanitize_filename_stem(name: str) -> str:
    for char in '\\/:*?"<>|':
        name = name.replace(char, "_")
    cleaned = name.strip().rstrip(".")
    return cleaned or "бриф"


def period_brief_display_name(
    period_start: date,
    period_end: date,
    brief_kind: str = "full",
) -> str:
    """Человекочитаемое имя брифа за период (без расширения)."""
    title = PERIOD_BRIEF_KIND_TITLES.get(brief_kind, PERIOD_BRIEF_KIND_TITLES["full"])
    start_fmt = period_start.strftime("%d.%m.%Y")
    end_fmt = period_end.strftime("%d.%m.%Y")
    if period_start == period_end:
        return f"{title} за {start_fmt}"
    return f"{title} {start_fmt}-{end_fmt}"


def weekly_brief_docx_filename(
    period_start: date,
    period_end: date,
    brief_kind: str = "full",
) -> str:
    stem = _sanitize_filename_stem(
        period_brief_display_name(period_start, period_end, brief_kind)
    )
    return f"{stem}.docx"


def period_brief_json_filename(
    period_start: date,
    period_end: date,
    brief_kind: str = "full",
) -> str:
    stem = _sanitize_filename_stem(
        period_brief_display_name(period_start, period_end, brief_kind)
    )
    return f"{stem}.json"


FIELD_LABELS_FOR_FORMATTING = {
    "заголовок новости",
    "продукт",
    "продукт/рынок",
    "рынок",
    "страна",
    "страна/регион",
    "регион",
    "суть меры",
    "ставка",
    "ставка/срок/охват",
    "срок",
    "охват",
    "источник",
    "источники",
    "тип проекта",
    "название проекта",
    "компания",
    "локация",
    "локация/мощность",
    "мощность",
    "capex",
    "сроки",
    "текущий статус",
    "событие",
    "метрики",
    "оценка результата",
    "ожидания",
    "ожидания/прогноз",
    "прогноз",
    "причины динамики",
    "комментарий аналитика",
    "комментарий аналитика/менеджмента",
    "кто покупает/продает",
    "объект сделки",
    "стоимость сделки",
    "прочая информация",
}

SECTION_ALIASES: list[tuple[str, str]] = [
    ("ежедневн", "title"),
    ("рыночный дайджест", "title"),
    ("динамика цен", "s1"),
    ("мировой рынок", "s1"),
    ("рынок стали и сырья в россии", "s1b"),
    ("экономики мира", "s2"),
    ("экономика россии", "s3"),
    ("балансы", "s4"),
    ("государственное регулирование", "s5"),
    ("кадровые", "s6"),
    ("проекты", "s7"),
    ("результаты компан", "s8"),
    ("m&a", "s9"),
    ("торговые барьер", "s10"),
    ("краткое резюме", "s11"),
]


@dataclass
class ParsedBrief:
    market_summary: str = ""
    products: dict[str, list[str]] = field(
        default_factory=lambda: {"ore": [], "coal": [], "scrap": []},
    )
    countries: dict[str, str] = field(default_factory=dict)
    russia: dict[str, str] = field(default_factory=dict)
    balance: list[str] = field(default_factory=list)
    regulation: list[str] = field(default_factory=list)
    hr: str = ""
    projects: str = ""
    results: str = ""
    ma: str = ""
    trade: str = ""
    summary: list[str] = field(default_factory=list)


def default_template_path() -> Path:
    return DEFAULT_TEMPLATE_PATH


def write_brief_docx(
    path: Path,
    body: str,
    *,
    title: str | None = None,
    report_date: str | None = None,
    period_label: str | None = None,
    metadata_lines: list[tuple[str, str]] | None = None,
    branded: bool = True,
    template_path: Path | None = None,
) -> Path:
    """Сохраняет готовый текст брифа в Word без заполнения шаблонных плейсхолдеров."""
    del title, metadata_lines, branded, template_path

    document = Document()
    _configure_prompt_document(document)
    _add_prompt_header(document)

    if report_date or period_label:
        info = document.add_paragraph()
        info.paragraph_format.space_after = Pt(10)
        run = info.add_run(
            " | ".join(
                part
                for part in (
                    f"Дата отчёта: {report_date}" if report_date else "",
                    f"Период: {period_label}" if period_label else "",
                )
                if part
            ),
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    current_section: str | None = None
    for line in body.splitlines():
        current_section = _add_prompt_line(document, line, current_section)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def _add_prompt_header(document: Document) -> None:
    logo_path = next((path for path in LOGO_CANDIDATES if path.is_file()), None)

    if logo_path:
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = logo_paragraph.add_run()
        run.add_picture(str(logo_path), width=Inches(2.1))

    disclaimer = document.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    disclaimer.paragraph_format.space_after = Pt(8)
    run = disclaimer.add_run("Подготовлено ИИ, может содержать ошибки")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _configure_prompt_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(4)

    for section in document.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.6)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)


def _add_prompt_line(document: Document, line: str, current_section: str | None = None) -> str | None:
    stripped = line.strip()
    if not stripped:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        return current_section
    if stripped in {"---", "— — —"}:
        return current_section

    numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
    if numbered and current_section in {"6", "7", "8", "9", "10"}:
        paragraph = document.add_paragraph()
        if int(numbered.group(1)) > 1:
            paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True
        _add_formatted_text(paragraph, stripped)
        return current_section

    if stripped.startswith("### "):
        paragraph = document.add_paragraph(style="Heading 3")
        _add_formatted_text(paragraph, stripped[4:].strip())
        return current_section
    if stripped.startswith("## "):
        heading_text = stripped[3:].strip()
        paragraph = document.add_paragraph(style="Heading 2")
        _add_formatted_text(paragraph, heading_text)
        return _section_number(heading_text)
    if stripped.startswith("# "):
        paragraph = document.add_paragraph(style="Heading 1")
        _add_formatted_text(paragraph, stripped[2:].strip())
        return current_section

    if _is_bullet(stripped):
        content = _clean_bullet(stripped)
        if _is_field_line(content) and not _field_should_be_bullet(content, current_section):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            _add_labeled_or_formatted_text(paragraph, content)
            return current_section

        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        _add_labeled_or_formatted_text(paragraph, content)
        return current_section

    if _is_field_line(stripped):
        paragraph = (
            document.add_paragraph(style="List Bullet")
            if _field_should_be_bullet(stripped, current_section)
            else document.add_paragraph()
        )
        paragraph.paragraph_format.space_after = Pt(3)
        _add_labeled_or_formatted_text(paragraph, stripped)
        return current_section

    paragraph = document.add_paragraph()
    if _is_template_subheading(stripped):
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(_strip_markdown_bold(stripped))
        run.bold = True
        run.font.size = Pt(10)
        return current_section

    if _is_case_title(stripped, current_section):
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(_strip_markdown_bold(stripped))
        run.bold = True
        run.font.size = Pt(10)
        return current_section

    _add_formatted_text(paragraph, stripped)
    return current_section


def _is_template_subheading(text: str) -> bool:
    clean = _strip_markdown_bold(text).rstrip(":")
    return (
        clean in TEMPLATE_PRODUCTS
        or clean in TEMPLATE_RUSSIA_MARKET
        or clean in TEMPLATE_COUNTRIES
        or clean in TEMPLATE_RUSSIA_TOPICS
        or clean.startswith("Динамика производства стали")
        or clean.startswith("Динамика экспорта стали")
        or clean.startswith("Динамика потребления стали")
    )


def _section_number(heading: str) -> str | None:
    match = re.match(r"^(\d+)\.", heading.strip())
    return match.group(1) if match else None


def _field_should_be_bullet(text: str, current_section: str | None) -> bool:
    del current_section
    return True


def _is_case_title(text: str, current_section: str | None) -> bool:
    if current_section not in {"6", "7", "8", "9", "10"}:
        return False
    if re.match(r"^\d+\.\s+", text.strip()):
        return False
    clean = _strip_markdown_bold(text).strip()
    if not clean or ":" in clean:
        return False
    if clean == NO_NEWS:
        return False
    return len(clean) <= 180


def _add_formatted_text(paragraph: Paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_labeled_or_formatted_text(paragraph: Paragraph, text: str) -> None:
    label, value = _split_field_label(text)
    if label is None:
        _add_formatted_text(paragraph, text)
        return

    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    _add_formatted_text(paragraph, value)


def _is_field_line(text: str) -> bool:
    label, _ = _split_field_label(text)
    return label is not None


def _split_field_label(text: str) -> tuple[str | None, str]:
    clean = _strip_markdown_bold(text).strip()
    if ":" not in clean:
        return None, clean
    label, value = clean.split(":", 1)
    normalized = label.strip().lower()
    if normalized not in FIELD_LABELS_FOR_FORMATTING:
        return None, clean
    return label.strip(), value.strip()


def convert_markdown_file_to_docx(
    markdown_path: Path,
    docx_path: Path | None = None,
    *,
    template_path: Path | None = None,
) -> Path:
    text = markdown_path.read_text(encoding="utf-8")
    if docx_path is None:
        docx_path = markdown_path.with_suffix(".docx")

    body = text
    report_date: str | None = None

    if text.startswith("#"):
        parts = text.split("\n---\n", 1)
        header_part = parts[0]
        body = parts[1] if len(parts) > 1 else ""
        for line in header_part.splitlines():
            line = line.strip()
            match = re.match(r"\*\*Сгенерировано:\*\*\s*(.+)", line)
            if match and not report_date:
                report_date = match.group(1).strip()[:10]
                if len(report_date) == 10 and report_date[4] == "-":
                    y, m, d = report_date.split("-")
                    report_date = f"{d}.{m}.{y}"

    write_brief_docx(
        docx_path,
        body.strip(),
        report_date=report_date,
        template_path=template_path,
    )
    return docx_path


def build_metadata_lines(period_range: Any, metadata: dict[str, Any]) -> list[tuple[str, str]]:
    """Служебные поля для markdown/manifest (не для шапки Word)."""
    return [
        ("Период", f"{period_range.start} — {period_range.end} ({period_range.name})"),
        ("Модель", str(metadata.get("model", ""))),
        ("Новостей в основе", str(metadata.get("news_count", 0))),
        ("Источников", str(metadata.get("sources_count", ""))),
        ("Сгенерировано", str(metadata.get("generated_at", ""))),
    ]


def parse_brief_body(body: str) -> ParsedBrief:
    """Разбирает markdown-текст LLM по разделам шаблона."""
    sections = _split_sections(body)
    parsed = ParsedBrief()

    parsed.market_summary = _join_non_bullets(sections.get("s1", []), stop_at_product=True)
    parsed.products = _extract_products(sections.get("s1", []))

    for country in TEMPLATE_COUNTRIES:
        parsed.countries[country] = _extract_labeled_block(sections.get("s2", []), country)

    for topic in TEMPLATE_RUSSIA_TOPICS:
        parsed.russia[topic] = _extract_labeled_block(sections.get("s3", []), topic)

    parsed.balance = _extract_balance(sections.get("s4", []))
    parsed.regulation = _collect_bullets(sections.get("s5", []))
    parsed.hr = _join_section_text(sections.get("s6", []))
    parsed.projects = _join_section_text(sections.get("s7", []))
    parsed.results = _join_section_text(sections.get("s8", []))
    parsed.ma = _join_section_text(sections.get("s9", []))
    parsed.trade = _join_section_text(sections.get("s10", []))
    parsed.summary = _collect_bullets(sections.get("s11", []))

    return parsed


def fill_docx_from_template(
    document: Document,
    parsed: ParsedBrief,
    report_date: str,
) -> None:
    """Заполняет копию шаблона, сохраняя стили и колонтитулы."""
    for paragraph in document.paragraphs:
        text = paragraph.text
        if "Ежедневный новостной дайджест" in text:
            _set_header_title(paragraph, report_date)
        if "КРАТКОЕ РЕЗЮМЕ" in text and "{ДАТА}" in text:
            _replace_in_paragraph(paragraph, "{ДАТА}", report_date)

    index = 0
    while index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        text = paragraph.text.strip()

        if _is_section_header(text, "1. МИРОВОЙ РЫНОК СТАЛИ И СЫРЬЯ"):
            index += 1
            index = _fill_next_placeholder(
                document,
                index,
                parsed.market_summary or NO_NEWS,
            )
            continue

        if text in TEMPLATE_PRODUCTS:
            key = PRODUCT_KEYS[text]
            title_index = index
            index = _fill_list_placeholders(
                document,
                index + 1,
                parsed.products.get(key) or [NO_NEWS],
                search_from=title_index + 1,
            )
            continue

        if text in TEMPLATE_COUNTRIES:
            index += 1
            index = _fill_next_placeholder(
                document,
                index,
                parsed.countries.get(text) or NO_NEWS,
            )
            continue

        if text in TEMPLATE_RUSSIA_TOPICS:
            index += 1
            index = _fill_next_placeholder(
                document,
                index,
                parsed.russia.get(text) or NO_NEWS,
            )
            continue

        if _is_section_header(text, "4. БАЛАНСЫ МИРОВОГО РЫНКА СТАЛИ"):
            index += 1
            balance = parsed.balance or [NO_NEWS, NO_NEWS, NO_NEWS]
            while len(balance) < 3:
                balance.append(NO_NEWS)
            for item in balance[:3]:
                index = _fill_next_placeholder(document, index, item)
            continue

        if _is_section_header(text, "5. ГОСУДАРСТВЕННОЕ РЕГУЛИРОВАНИЕ"):
            index += 1
            index = _fill_list_placeholders(
                document,
                index,
                parsed.regulation or [NO_NEWS],
            )
            continue

        if _is_section_header(text, "6. КАДРОВЫЕ НАЗНАЧЕНИЯ"):
            index += 1
            index = _fill_block_section(document, index, parsed.hr or NO_NEWS)
            continue

        if _is_section_header(text, "7. ПРОЕКТЫ ПО СТРОИТЕЛЬСТВУ"):
            index += 1
            index = _fill_block_section(document, index, parsed.projects or NO_NEWS)
            continue

        if _is_section_header(text, "8. РЕЗУЛЬТАТЫ КОМПАНИЙ"):
            index += 1
            index = _fill_block_section(document, index, parsed.results or NO_NEWS)
            continue

        if _is_section_header(text, "9. M&A СДЕЛКИ"):
            index += 1
            index = _fill_block_section(document, index, parsed.ma or NO_NEWS)
            continue

        if _is_section_header(text, "10. ТОРГОВЫЕ БАРЬЕРЫ"):
            index += 1
            index = _fill_block_section(document, index, parsed.trade or NO_NEWS)
            continue

        if _is_section_header(text, "11. КРАТКОЕ РЕЗЮМЕ"):
            index += 1
            index = _fill_list_placeholders(
                document,
                index,
                parsed.summary or [NO_NEWS],
            )
            continue

        if _is_instruction_line(text):
            _remove_paragraph(paragraph)
            continue

        index += 1


# --- разбор markdown ---


def _split_sections(body: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = "_intro"

    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line or line in {"---", "— — —"}:
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            current = _section_key(heading.group(2))
            sections.setdefault(current, [])
            continue

        numbered = re.match(r"^(\d+)\)\s+(.+)$", line)
        if numbered:
            current = _section_key(numbered.group(2))
            sections.setdefault(current, [])
            continue

        sections.setdefault(current, []).append(line)

    return sections


def _section_key(title: str) -> str:
    lowered = _strip_markdown_bold(title).lower()
    for needle, key in SECTION_ALIASES:
        if needle in lowered:
            return key
    return "_other"


def _extract_products(lines: list[str]) -> dict[str, list[str]]:
    products: dict[str, list[str]] = {"ore": [], "coal": [], "scrap": []}
    current: str | None = None
    summary_bits: list[str] = []

    for line in lines:
        if _is_bullet(line):
            text = _clean_bullet(line)
            if current:
                products[current].append(text)
            else:
                summary_bits.append(text)
            continue

        product = _match_product_key(line)
        if product:
            current = product
            continue

        if not current:
            summary_bits.append(_strip_markdown_bold(line))

    if summary_bits and not any(products.values()):
        pass  # summary handled separately
    return products


def _join_non_bullets(lines: list[str], *, stop_at_product: bool = False) -> str:
    chunks: list[str] = []
    for line in lines:
        if stop_at_product and _match_product_key(line):
            break
        if _is_bullet(line):
            continue
        chunks.append(_strip_markdown_bold(line))
    return " ".join(chunk for chunk in chunks if chunk).strip()


def _extract_labeled_block(lines: list[str], label: str) -> str:
    label_lower = label.lower()
    for index, line in enumerate(lines):
        plain = _strip_markdown_bold(line)
        if plain.lower().startswith(label_lower):
            value = plain[len(label) :].lstrip(" :—-").strip()
            if value:
                return value
            following = _next_non_bullet_paragraph(lines, index + 1)
            return following or NO_NEWS
    return ""


def _next_non_bullet_paragraph(lines: list[str], start: int) -> str:
    for line in lines[start:]:
        if _is_bullet(line):
            continue
        if _match_product_key(line) or any(
            line.strip().startswith(country) for country in TEMPLATE_COUNTRIES
        ):
            break
        text = _strip_markdown_bold(line)
        if text:
            return text
    return ""


def _extract_balance(lines: list[str]) -> list[str]:
    buckets: dict[str, list[str]] = {key: [] for key in BALANCE_LABELS}

    for line in lines:
        text = _clean_bullet(line) if _is_bullet(line) else _strip_markdown_bold(line)
        if not text:
            continue
        lowered = text.lower()
        for key in BALANCE_LABELS:
            if key in lowered:
                buckets[key].append(text)
                break

    result = [" ".join(buckets[key]).strip() for key in BALANCE_LABELS]
    if not any(result):
        bullets = [_clean_bullet(line) for line in lines if _is_bullet(line)]
        if bullets:
            return bullets[:3]
        paragraphs = [
            _strip_markdown_bold(line)
            for line in lines
            if line and not _is_bullet(line)
        ]
        return paragraphs[:3]
    return [item or NO_NEWS for item in result]


def _collect_bullets(lines: list[str]) -> list[str]:
    items = [_clean_bullet(line) for line in lines if _is_bullet(line)]
    if items:
        return items
    text = _join_section_text(lines)
    return [text] if text else []


def _join_section_text(lines: list[str]) -> str:
    chunks: list[str] = []
    for line in lines:
        if _is_instruction_line(line):
            continue
        if _is_bullet(line):
            chunks.append(_clean_bullet(line))
        else:
            text = _strip_markdown_bold(line)
            if text:
                chunks.append(text)
    return "\n".join(chunks).strip()


def _match_product_key(line: str) -> str | None:
    lowered = _strip_markdown_bold(line).lower()
    if "железн" in lowered and "руд" in lowered:
        return "ore"
    if "кокс" in lowered or "коксующ" in lowered:
        return "coal"
    if "лом" in lowered and ("турц" in lowered or "turkey" in lowered or "рынок" in lowered):
        return "scrap"
    if lowered in {
        "железная руда в китае",
        "коксующийся уголь",
        "стальной лом в турции",
    }:
        return {"железная руда в китае": "ore", "коксующийся уголь": "coal", "стальной лом в турции": "scrap"}[
            lowered
        ]
    return None


# --- заполнение шаблона ---


def _fill_next_placeholder(
    document: Document,
    index: int,
    text: str,
) -> int:
    while index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        if _is_placeholder(paragraph.text):
            _set_paragraph_text(paragraph, text)
            return index + 1
        if not paragraph.text.strip():
            index += 1
            continue
        break
    return index


def _fill_list_placeholders(
    document: Document,
    index: int,
    items: list[str],
    *,
    search_from: int | None = None,
) -> int:
    placeholders: list[Paragraph] = []
    while index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        text = paragraph.text.strip()
        if _is_placeholder(text) and paragraph.style.name == "List Paragraph":
            placeholders.append(paragraph)
            index += 1
            continue
        if not text:
            index += 1
            continue
        break

    if not placeholders:
        return index

    for position, item in enumerate(items):
        if position < len(placeholders):
            _set_paragraph_text(placeholders[position], item)
        else:
            clone = _clone_paragraph_after(placeholders[-1])
            _set_paragraph_text(clone, item)

    for extra in placeholders[len(items) :]:
        _remove_paragraph(extra)

    return _find_next_anchor_index(document, search_from if search_from is not None else index)


def _find_next_anchor_index(document: Document, index: int) -> int:
    """Индекс следующего заголовка продукта, страны, темы России или раздела."""
    while index < len(document.paragraphs):
        stripped = document.paragraphs[index].text.strip()
        if not stripped:
            index += 1
            continue
        if (
            stripped in TEMPLATE_PRODUCTS
            or stripped in TEMPLATE_COUNTRIES
            or stripped in TEMPLATE_RUSSIA_TOPICS
            or _is_section_header(stripped, "")
        ):
            return index
        index += 1
    return index


def _fill_block_section(document: Document, index: int, text: str) -> int:
    while index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        stripped = paragraph.text.strip()
        if _is_section_header(stripped, ""):
            break
        if _is_instruction_line(stripped):
            _remove_paragraph(paragraph)
            continue
        if _is_placeholder(stripped):
            _set_paragraph_text(paragraph, text)
            index += 1
            continue
        if stripped.startswith("Типы проектов:"):
            index += 1
            continue
        break
    return index


def _set_header_title(paragraph: Paragraph, report_date: str) -> None:
    new_text = f"Ежедневный новостной дайджест ({report_date})"
    if paragraph.text.strip().startswith("Ежедневный"):
        _set_paragraph_text(paragraph, new_text, preserve_runs=False)


def _set_paragraph_text(paragraph: Paragraph, text: str, *, preserve_runs: bool = True) -> None:
    if paragraph.runs:
        if preserve_runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            for run in list(paragraph.runs):
                run.text = ""
            paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> None:
    if old in paragraph.text:
        _set_paragraph_text(paragraph, paragraph.text.replace(old, new))


def _clone_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_element = deepcopy(paragraph._element)
    paragraph._element.addnext(new_element)
    return Paragraph(new_element, paragraph._parent)


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _is_placeholder(text: str) -> bool:
    stripped = text.strip()
    return stripped.startswith("{") or stripped.startswith(" {") or "{" in stripped[:3]


def _is_instruction_line(text: str) -> bool:
    lowered = text.strip().lower()
    return lowered.startswith("если нет данных") or lowered.startswith("типы проектов:")


def _is_section_header(text: str, marker: str) -> bool:
    if marker and marker in text:
        return True
    return bool(re.match(r"^\d{1,2}\.\s+[А-ЯA-Z]", text))


def _is_bullet(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("* ")


def _clean_bullet(line: str) -> str:
    return _strip_markdown_bold(line.lstrip("-•* ").strip())


def _strip_markdown_bold(text: str) -> str:
    return re.sub(r"\*\*", "", text).strip()
